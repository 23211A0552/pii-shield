import io
import logging
import fitz  # PyMuPDF
from docx import Document
from PIL import Image, ImageDraw, ImageOps
from typing import List, Dict, Any

logger = logging.getLogger(__name__)

def redact_pdf(pdf_bytes: bytes, detections: List[Dict[str, Any]]) -> bytes:
    """
    Redact a PDF by drawing black boxes over detected PII.
    Preserves the original document's visual structure across ALL pages.
    Correctly collects all annotations per page before applying redactions.
    """
    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")

        for page in doc:
            # Collect ALL redact annotations for this page first,
            # then apply once — calling apply_redactions per annotation
            # can cause missed redactions in PyMuPDF.
            found_any = False
            for detection in detections:
                pii_text = detection.get("detected_value")
                if not pii_text:
                    continue
                text_instances = page.search_for(pii_text, quads=True)
                for quad in text_instances:
                    page.add_redact_annot(quad, fill=(0, 0, 0))
                    found_any = True

            if found_any:
                page.apply_redactions(images=fitz.PDF_REDACT_IMAGE_NONE)

        output_buffer = io.BytesIO()
        doc.save(output_buffer, garbage=4, deflate=True)
        doc.close()
        return output_buffer.getvalue()

    except Exception as e:
        logger.error(f"PDF redaction failed: {e}")
        return pdf_bytes

def redact_docx(docx_bytes: bytes, detections: List[Dict[str, Any]]) -> bytes:
    """
    Redact a Word document by replacing PII text in runs.
    Handles PII that is split across multiple runs (bold/italic boundaries).
    """
    try:
        doc = Document(io.BytesIO(docx_bytes))

        # Sorted by length descending so longer matches take priority
        to_redact = sorted(
            {d["detected_value"] for d in detections if d.get("detected_value")},
            key=len, reverse=True
        )

        def _redact_paragraph(para) -> None:
            """
            Redact a paragraph, handling PII that spans run boundaries.
            Strategy:
              1. Check if the full paragraph text contains any PII.
              2. If yes, do replacement at text level, then redistribute back
                 into runs (clearing overflow runs).
              3. If runs contain PII individually, do simple per-run replacement.
            """
            combined = para.text
            needs_redact = any(val in combined for val in to_redact)
            if not needs_redact:
                return

            # Apply all replacements to the combined text
            redacted = combined
            for val in to_redact:
                redacted = redacted.replace(val, "[REDACTED]")

            # Now redistribute the redacted text back into runs
            # We put all text into the FIRST run and blank the rest
            # (preserving formatting of first run as best-effort)
            if para.runs:
                para.runs[0].text = redacted
                for run in para.runs[1:]:
                    run.text = ""
            else:
                # No runs — set the paragraph XML text node directly
                for child in para._p:
                    if child.text is not None:
                        child.text = redacted
                        break

        # Process all paragraphs
        for para in doc.paragraphs:
            _redact_paragraph(para)

        # Process tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        _redact_paragraph(para)

        output_buffer = io.BytesIO()
        doc.save(output_buffer)
        return output_buffer.getvalue()

    except Exception as e:
        logger.error(f"DOCX redaction failed: {e}")
        return docx_bytes


def _ext_to_pil_format(ext: str) -> str:
    """
    Map a file extension to the Pillow save format string.
    Always returns a format that Pillow can reliably save.
    """
    ext = ext.lower().lstrip(".")
    mapping = {
        "jpg":  "JPEG",
        "jpeg": "JPEG",
        "jfif": "JPEG",
        "jpe":  "JPEG",
        "png":  "PNG",
        "bmp":  "BMP",
        "gif":  "GIF",
        "tiff": "TIFF",
        "tif":  "TIFF",
        "webp": "WEBP",
    }
    return mapping.get(ext, "JPEG")  # default to JPEG for unknown image types


def redact_image(
    image_bytes: bytes,
    detections: List[Dict[str, Any]],
    scale_factor: float = 1.0,
    filename: str = "",
) -> bytes:
    """
    Redact an image by drawing solid black boxes over detected PII.

    Args:
        image_bytes:   Original image bytes (the file that will be downloaded).
        detections:    PII detections with x/y/w/h bbox from Tesseract.
        scale_factor:  Ratio original_width / ocr_width.  Pass this when
                       Tesseract ran on an upscaled image so coordinates are
                       mapped back to original pixel space.
        filename:      Original filename — used to determine the output format
                       reliably (avoids MPO/None issues from img.format).
    """
    try:
        img = Image.open(io.BytesIO(image_bytes))
        img = ImageOps.exif_transpose(img)  # Standardise orientation

        # Determine output format from the file extension (most reliable source).
        # Pillow's img.format can be "MPO" for camera JPEGs, which Windows
        # image viewers reject. We always normalise to a plain saveable format.
        if filename:
            ext = filename.rsplit(".", 1)[-1] if "." in filename else "jpg"
            save_fmt = _ext_to_pil_format(ext)
        else:
            # Fallback: use Pillow's detected format, map JPEG variants → JPEG
            detected = (img.format or "JPEG").upper()
            save_fmt = "JPEG" if detected in ("JPEG", "MPO", "JPEG2000", "JPX", "JPG") else detected
            save_fmt = "JPEG" if save_fmt not in ("PNG", "BMP", "GIF", "TIFF", "WEBP", "JPEG") else save_fmt

        # Convert colour mode so Pillow can save cleanly.
        # JPEG cannot handle RGBA or CMYK → force RGB.
        if save_fmt == "JPEG":
            if img.mode in ("RGBA", "P", "LA"):
                img = img.convert("RGB")
            elif img.mode == "CMYK":
                img = img.convert("RGB")
            elif img.mode != "RGB":
                img = img.convert("RGB")
        elif img.mode not in ("RGB", "RGBA", "L"):
            img = img.convert("RGB")

        draw = ImageDraw.Draw(img)

        redacted_any = False
        for det in detections:
            # Use the full bboxes list (all occurrences) when available,
            # otherwise fall back to the single x/y/w/h fields.
            bboxes = det.get("bboxes")
            if not bboxes:
                x = det.get("x")
                y = det.get("y")
                w = det.get("w")
                h = det.get("h")
                if None in (x, y, w, h):
                    continue
                bboxes = [{"x": x, "y": y, "w": w, "h": h}]

            for bbox in bboxes:
                x = bbox.get("x")
                y = bbox.get("y")
                w = bbox.get("w")
                h = bbox.get("h")
                if None in (x, y, w, h):
                    continue

                # Scale coordinates back to original image space
                ox = int(x * scale_factor)
                oy = int(y * scale_factor)
                ow = int(w * scale_factor)
                oh = int(h * scale_factor)

                padding = max(2, int(4 * scale_factor))  # proportional padding
                draw.rectangle(
                    [ox - padding, oy - padding, ox + ow + padding, oy + oh + padding],
                    fill="black"
                )
                redacted_any = True

        if not redacted_any:
            logger.warning("Image redaction: no coordinates found — returning original.")

        output_buffer = io.BytesIO()
        save_kwargs: Dict[str, Any] = {"format": save_fmt}
        if save_fmt == "JPEG":
            save_kwargs["quality"] = 95       # high quality, no visible artefacts
            save_kwargs["subsampling"] = 0    # 4:4:4 chroma — best quality
        img.save(output_buffer, **save_kwargs)
        logger.info("Image redacted and saved as %s (scale=%.3f)", save_fmt, scale_factor)
        return output_buffer.getvalue()

    except Exception as e:
        logger.error(f"Image redaction failed: {e}")
        return image_bytes