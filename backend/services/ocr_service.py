"""
ocr_extractor.py
~~~~~~~~~~~~~~~~
Text extraction from images, PDFs, and DOCX files.

Improvements:
- Multi-pass Tesseract OCR (PSM 3, 6, 4, 11) — picks best result
- Adaptive preprocessing (sharpen, contrast, brightness, upscale)
- Document-type hint → bilingual OCR for Aadhaar (eng+hin)
- PyMuPDF → PyPDF2 → pdf2image+OCR fallback chain for PDFs
- Table extraction from DOCX (catches structured PII in forms)
"""

import io
import logging
import os
from pathlib import Path
from typing import Any, List, Dict

import pytesseract
from PIL import Image, ImageFilter, ImageOps, ImageEnhance

logger = logging.getLogger(__name__)

# Optional: override Tesseract binary path via env var
_tess_path = os.getenv("TESSERACT_PATH")
if _tess_path:
    pytesseract.pytesseract.tesseract_cmd = _tess_path


# ══════════════════════════════════════════════════════════════════════════════
# IMAGE PREPROCESSING
# ══════════════════════════════════════════════════════════════════════════════

def _preprocess(image: Image.Image) -> Image.Image:
    """
    Standard preprocessing pipeline for ID-card / scanned-document OCR.
    Returns a grayscale Pillow image ready for Tesseract.
    """
    img = ImageOps.grayscale(image)

    # Upscale small images (phone snapshots of cards are often low-res)
    if img.width < 1500:
        scale = 2000 / img.width
        img = img.resize(
            (int(img.width * scale), int(img.height * scale)),
            Image.Resampling.LANCZOS,
        )

    img = img.filter(ImageFilter.SHARPEN)
    img = ImageEnhance.Contrast(img).enhance(2.0)
    img = ImageEnhance.Brightness(img).enhance(1.1)
    return img


# ══════════════════════════════════════════════════════════════════════════════
# MULTI-PASS OCR
# ══════════════════════════════════════════════════════════════════════════════

def _ocr_multi_pass(image: Image.Image, lang: str = "eng") -> tuple:
    """Try multiple Tesseract PSM modes and return the text with most content."""
    best_text = ""
    best_psm = 3
    for psm in (3, 6, 4, 11):
        try:
            config = f"--psm {psm} --oem 3"
            text = pytesseract.image_to_string(image, lang=lang, config=config)
            if len(text.strip()) > len(best_text.strip()):
                best_text = text
                best_psm = psm
        except Exception as exc:
            logger.debug("OCR PSM %d failed: %s", psm, exc)
    return best_text.strip(), best_psm


# ══════════════════════════════════════════════════════════════════════════════
# PUBLIC EXTRACTION FUNCTIONS
# ══════════════════════════════════════════════════════════════════════════════

def _hint_to_lang(hint: str) -> str:
    """Map document hint to Tesseract language string."""
    if hint in ("aadhaar",):
        # Try Hindi+English for Aadhaar; fall back to eng if hin not installed
        try:
            pytesseract.get_languages(config="")
            langs = pytesseract.get_languages(config="")
            if "hin" in langs:
                return "eng+hin"
        except Exception:
            pass
    return "eng"


def _get_image_frames(image_bytes: bytes) -> List[Image.Image]:
    """
    Return a list of Pillow Image frames.
    Handles multi-page TIFFs; single-frame images return a 1-element list.
    """
    frames = []
    try:
        img = Image.open(io.BytesIO(image_bytes))
        while True:
            frames.append(img.copy())
            try:
                img.seek(img.tell() + 1)
            except EOFError:
                break
    except Exception as exc:
        logger.error("Failed to read image frames: %s", exc)
    return frames if frames else []


def extract_text_from_image(image_bytes: bytes, hint: str = "general", return_coords: bool = False) -> Any:
    """
    Extract text (and optionally coordinates) from an image.
    Supports multi-page / multi-frame TIFFs — all frames are processed.
    If return_coords is True, returns a list of word dicts with bbox info.
    """
    try:
        frames = _get_image_frames(image_bytes)
        if not frames:
            return [] if return_coords else ""

        lang = _hint_to_lang(hint)
        all_words: List[Dict] = []
        all_texts: List[str] = []

        for frame_idx, image in enumerate(frames):
            # 1. Handle EXIF orientation (ensure phone photos are right-side up)
            image = ImageOps.exif_transpose(image)
            
            processed = _preprocess(image)

            # Multi-pass: find best PSM for this frame
            try:
                text, psm = _ocr_multi_pass(processed, lang=lang)
            except Exception:
                text, psm = _ocr_multi_pass(processed, lang="eng")
                lang = "eng"

            if not return_coords:
                all_texts.append(text)
                continue

            # Get word-level bounding boxes using the best PSM
            config = f"--psm {psm} --oem 3"
            data = pytesseract.image_to_data(
                processed, lang=lang, config=config,
                output_type=pytesseract.Output.DICT
            )
            for i in range(len(data["text"])):
                conf = float(data["conf"][i])
                word_text = data["text"][i].strip()
                if conf > 10 and word_text:
                    all_words.append({
                        "text": word_text,
                        "x": data["left"][i],
                        "y": data["top"][i],
                        "w": data["width"][i],
                        "h": data["height"][i],
                        "conf": conf,
                        "line_num": data["line_num"][i],
                        "block_num": data["block_num"][i],
                        "page": frame_idx,   # track which frame this word came from
                    })

        if return_coords:
            return all_words
        return "\n\n--- Page Break ---\n\n".join(all_texts)

    except Exception as exc:
        logger.error("Image OCR failed: %s", exc)
        return [] if return_coords else ""


def extract_text_from_pdf(pdf_bytes: bytes) -> str:
    """
    Extract text from a PDF using a three-tier fallback strategy:
      1. PyMuPDF (fitz)  — best for digital PDFs with text layer
      2. PyPDF2           — pure-Python fallback
      3. pdf2image + OCR  — last resort for scanned/image-only PDFs
    """
    # ── Tier 1: PyMuPDF ──────────────────────────────────────────────────────
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        parts = [page.get_text("text") for page in doc]
        text = "\n".join(parts).strip()
        if text:
            return text
        logger.info("PyMuPDF: no text layer – falling back to OCR.")
    except ImportError:
        logger.debug("PyMuPDF not installed.")
    except Exception as exc:
        logger.warning("PyMuPDF error (%s) – trying PyPDF2.", exc)

    # ── Tier 2: PyPDF2 ───────────────────────────────────────────────────────
    try:
        import PyPDF2
        reader = PyPDF2.PdfReader(io.BytesIO(pdf_bytes))
        parts = [p.extract_text() or "" for p in reader.pages]
        text = "\n".join(parts).strip()
        if text:
            return text
    except ImportError:
        logger.debug("PyPDF2 not installed.")
    except Exception as exc:
        logger.warning("PyPDF2 error (%s) – falling back to OCR.", exc)

    # ── Tier 3: OCR ──────────────────────────────────────────────────────────
    return _ocr_pdf(pdf_bytes)


def _ocr_pdf(pdf_bytes: bytes) -> str:
    """Convert each PDF page to an image and run OCR."""
    try:
        from pdf2image import convert_from_bytes
        images = convert_from_bytes(pdf_bytes, dpi=250)
        parts = []
        for img in images:
            buf = io.BytesIO()
            img.save(buf, format="PNG")
            parts.append(extract_text_from_image(buf.getvalue()))
        return "\n".join(parts)
    except Exception as exc:
        logger.error("PDF OCR failed: %s", exc)
        raise ValueError(f"Failed to OCR PDF: {exc}") from exc


def extract_text_from_docx(docx_bytes: bytes) -> str:
    """
    Extract text from a DOCX file including both paragraphs and tables.
    Tables often contain structured PII in government / bank forms.
    """
    try:
        from docx import Document
        doc = Document(io.BytesIO(docx_bytes))
        parts: list[str] = []

        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())

        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))

        return "\n".join(parts)

    except Exception as exc:
        logger.error("DOCX extraction failed: %s", exc)
        raise ValueError(f"Failed to extract text from DOCX: {exc}") from exc


# ══════════════════════════════════════════════════════════════════════════════
# SMART ROUTER
# ══════════════════════════════════════════════════════════════════════════════

def _guess_hint(filename: str) -> str:
    """Infer document type from filename to optimise OCR settings."""
    name = filename.lower()
    if any(k in name for k in ("aadhaar", "aadhar", "uid")):
        return "aadhaar"
    if any(k in name for k in ("pan", "income_tax")):
        return "pan"
    if "passport" in name:
        return "passport"
    if any(k in name for k in ("credit", "debit", "card", "visa", "mastercard")):
        return "credit_card"
    if any(k in name for k in ("bank", "statement", "passbook", "cheque", "check")):
        return "bank"
    return "general"


def extract_text(file_bytes: bytes, content_type: str, filename: str) -> str:
    """
    Route file bytes to the correct extractor based on MIME type / extension.
    Returns extracted plain text (line-structured for PII parsing).
    """
    ext = Path(filename).suffix.lower()
    ct  = content_type.lower()

    if ct.startswith("image/") or ext in (".jpg", ".jpeg", ".png", ".bmp",
                                           ".tiff", ".tif", ".gif", ".webp"):
        hint = _guess_hint(filename)
        # Build line-structured text from word-level data (preserves line boundaries)
        words = extract_text_from_image(file_bytes, hint=hint, return_coords=True)
        text = _words_to_structured_text(words)
        return _clean_extracted_text(text)

    elif ct == "application/pdf" or ext == ".pdf":
        return _clean_extracted_text(extract_text_from_pdf(file_bytes))

    elif ext == ".docx" or "wordprocessingml" in ct or "msword" in ct:
        return _clean_extracted_text(extract_text_from_docx(file_bytes))

    elif ext == ".txt" or ct.startswith("text/"):
        return _clean_extracted_text(file_bytes.decode("utf-8", errors="replace"))

    else:
        try:
            decoded = file_bytes.decode("utf-8", errors="replace")
            if decoded.strip():
                return _clean_extracted_text(decoded)
        except Exception:
            pass
        raise ValueError(f"Unsupported file type: {content_type} / {ext}")


def extract_text_with_coords(file_bytes: bytes, content_type: str, filename: str):
    """
    For images: returns (plain_text, word_map) where word_map is a list of
    word dicts with bbox coordinates for use during redaction.
    For non-image files, returns (plain_text, []).
    """
    ext = Path(filename).suffix.lower()
    ct  = content_type.lower()

    if ct.startswith("image/") or ext in (".jpg", ".jpeg", ".png", ".bmp",
                                           ".tiff", ".tif", ".gif", ".webp"):
        hint = _guess_hint(filename)
        words = extract_text_from_image(file_bytes, hint=hint, return_coords=True)
        text = _clean_extracted_text(_words_to_structured_text(words))
        return text, words

    # For PDF / DOCX / text, return text only (no pixel coords available)
    text = extract_text(file_bytes, content_type, filename)
    return text, []


def _words_to_structured_text(words: List[Dict]) -> str:
    """
    Reconstruct line-by-line text from Tesseract word dicts,
    grouping by (block_num, line_num) to preserve structure.
    """
    if not words:
        return ""
    from collections import defaultdict
    lines: Dict[tuple, list] = defaultdict(list)
    for w in words:
        key = (w.get("block_num", 0), w.get("line_num", 0))
        lines[key].append(w["text"])
    result = []
    for key in sorted(lines.keys()):
        result.append(" ".join(lines[key]))
    return "\n".join(result)

def _clean_extracted_text(text: str) -> str:
    """
    Light-touch cleanup: normalize whitespace but PRESERVE individual lines
    so that structured parsers (PAN, Aadhaar) can match label+value patterns.
    """
    import re

    if not text:
        return ""

    # Normalize line endings
    text = text.replace("\r\n", "\n").replace("\r", "\n")

    cleaned_lines = []
    for line in text.split("\n"):
        stripped = line.strip()
        # Drop lines that are pure noise (only special chars, length ≤ 2)
        if len(stripped) <= 2 and not any(c.isalnum() for c in stripped):
            continue
        # Collapse internal whitespace runs
        stripped = re.sub(r" +", " ", stripped)
        cleaned_lines.append(stripped)

    # Remove consecutive blank lines (allow max 1 blank between sections)
    final_lines = []
    prev_blank = False
    for line in cleaned_lines:
        is_blank = not line
        if is_blank and prev_blank:
            continue
        final_lines.append(line)
        prev_blank = is_blank

    return "\n".join(final_lines).strip()